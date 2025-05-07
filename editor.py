import tkinter as tk
from tkinter import ttk, font, filedialog, messagebox, colorchooser
from tkinter.scrolledtext import ScrolledText
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
import threading
from PIL import Image, ImageTk
from io import BytesIO
import speech_recognition as sr
import time
from ttkthemes import ThemedTk
from gtts import gTTS
import tempfile
import pygame
import fitz  # PyMuPDF for PDF export
import pytesseract
from langdetect import detect
from deep_translator import GoogleTranslator
import re

# Configure Tesseract OCR path
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

class DocumentEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Editor")
        self.root.geometry("1200x800")
        
        # Document state
        self.current_file = None
        self.is_modified = False
        self.current_theme = "light"
        self.autosave_interval = 300  # 5 minutes
        self.current_audio = None  # Track current audio file
        self.voice_typing_active = False  # Initialize voice typing state
        self.ocr_cancel_flag = False  # Add this line
        
        # Initialize pygame mixer for audio
        pygame.mixer.init()
        self.is_reading = False
        
        self.setup_ui()
        self.setup_bindings()
        self.start_autosave()
        self.load_settings()
        
    def setup_ui(self):
        # Configure modern style for the editor
        style = ttk.Style()
        style.configure('Editor.TFrame', background='#f0f0f0')
        style.configure('Ribbon.TFrame', background='#f8f9fa')
        style.configure('Ribbon.TButton', padding=5)  # Increased padding for better touch targets
        style.configure('Tool.TButton', padding=3, width=10)  # Fixed width for tool buttons
        
        # Create main container
        self.main_container = ttk.Frame(self.root, style='Editor.TFrame')
        self.main_container.pack(fill=tk.BOTH, expand=True)
        
        # Create ribbon with modern styling
        self.create_ribbon()
        
        # Add separator between ribbon and text area
        ttk.Separator(self.main_container, orient='horizontal').pack(fill=tk.X, pady=(0, 2))
        
        # Main text area
        self.create_text_area()
        
        # Status bar
        self.create_status_bar()
        
    def create_ribbon(self):
        # Create main ribbon frame
        ribbon_frame = ttk.Frame(self.main_container, style='Ribbon.TFrame')
        ribbon_frame.pack(fill=tk.X, padx=2, pady=2)
        
        # File Operations Group
        file_frame = ttk.LabelFrame(ribbon_frame, text="File", padding=(5, 5))
        file_frame.pack(side=tk.LEFT, padx=5, pady=2)
        
        file_buttons = [
            ("New", "📄 New", self.new_document),
            ("Open", "📂 Open", self.open_document),
            ("Save", "💾 Save", self.save_document),
            ("Save As", "💾 Save As", self.save_as_document)
        ]
        
        for tooltip, text, command in file_buttons:
            btn = ttk.Button(file_frame, text=text, command=command, style='Tool.TButton')
            btn.pack(side=tk.LEFT, padx=2, pady=2)
            self.create_tooltip(btn, tooltip)
        
        # Format Group
        format_frame = ttk.LabelFrame(ribbon_frame, text="Format", padding=(5, 5))
        format_frame.pack(side=tk.LEFT, padx=5, pady=2)
        
        # Font controls with better layout
        font_frame = ttk.Frame(format_frame)
        font_frame.pack(side=tk.LEFT, padx=2)
        
        ttk.Label(font_frame, text="Font:").pack(side=tk.LEFT)
        self.font_var = tk.StringVar(value="Consolas")
        font_combo = ttk.Combobox(font_frame, textvariable=self.font_var, width=15)
        font_combo['values'] = sorted(['Consolas', 'Courier New', 'Roboto Mono'] + 
                                    [f for f in font.families() if 'mono' in f.lower()])
        font_combo.pack(side=tk.LEFT, padx=2)
        font_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_font())
        
        ttk.Label(font_frame, text="Size:").pack(side=tk.LEFT, padx=(5, 0))
        self.size_var = tk.StringVar(value="12")
        size_combo = ttk.Combobox(font_frame, textvariable=self.size_var, width=4)
        size_combo['values'] = [8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24]
        size_combo.pack(side=tk.LEFT, padx=2)
        size_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_font())
        
        ttk.Button(format_frame, text="🎨 Color", command=self.choose_color, 
                   style='Tool.TButton').pack(side=tk.LEFT, padx=2)
        
        # Tools Group
        tools_frame = ttk.LabelFrame(ribbon_frame, text="Tools", padding=(5, 5))
        tools_frame.pack(side=tk.LEFT, padx=5, pady=2)
        
        # Voice tools with icons
        self.voice_button = ttk.Button(tools_frame, text="🎤 Voice", 
                                      command=self.toggle_voice_typing, style='Tool.TButton')
        self.voice_button.pack(side=tk.LEFT, padx=2)
        self.create_tooltip(self.voice_button, "Start/Stop Voice Typing")
        
        self.read_button = ttk.Button(tools_frame, text="🔊 Read", 
                                     command=self.read_text, style='Tool.TButton')
        self.read_button.pack(side=tk.LEFT, padx=2)
        self.create_tooltip(self.read_button, "Read Text Aloud")
        
        ocr_button = ttk.Button(tools_frame, text="📷 OCR", 
                                 command=self.show_ocr_dialog, style='Tool.TButton')
        ocr_button.pack(side=tk.LEFT, padx=2)
        self.create_tooltip(ocr_button, "Extract Text from Images/PDFs")
        
        # Theme toggle with icon
        theme_button = ttk.Button(tools_frame, text="🌓 Theme", 
                                 command=self.toggle_theme, style='Tool.TButton')
        theme_button.pack(side=tk.LEFT, padx=2)
        self.create_tooltip(theme_button, "Toggle Light/Dark Theme")

    def create_tooltip(self, widget, text):
        """Create a tooltip for a given widget"""
        def show_tooltip(event):
            tooltip = tk.Toplevel()
            tooltip.wm_overrideredirect(True)
            tooltip.wm_geometry(f"+{event.x_root+10}+{event.y_root+10}")
            
            label = ttk.Label(tooltip, text=text, background="#ffffe0", 
                             relief='solid', borderwidth=1)
            label.pack()
            
            def hide_tooltip():
                tooltip.destroy()
            
            widget.tooltip = tooltip
            widget.bind('<Leave>', lambda e: hide_tooltip())
            tooltip.bind('<Leave>', lambda e: hide_tooltip())
        
        widget.bind('<Enter>', show_tooltip)

    def create_text_area(self):
        text_frame = ttk.Frame(self.main_container)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        
        self.text_area = ScrolledText(text_frame, wrap=tk.WORD, undo=True)
        self.text_area.pack(fill=tk.BOTH, expand=True)
        
        # Configure modern text editor styling
        self.text_area.configure(
            font=("Consolas", 12),
            bg="#ffffff",
            fg="#212529",
            insertbackground="#212529",
            selectbackground="#c2dbff",
            selectforeground="#212529",
            insertwidth=2,
            spacing1=0,
            spacing2=0,
            spacing3=0,
            padx=5,
            pady=5
        )
        
    def create_status_bar(self):
        status = ttk.Frame(self.main_container)
        status.pack(fill=tk.X, side=tk.BOTTOM)
        
        # Add top border to status bar
        separator = ttk.Separator(self.main_container, orient='horizontal')
        separator.pack(fill=tk.X, side=tk.BOTTOM)
        
        # Modern status bar styling
        self.status_label = ttk.Label(status, text="Ready", padding=(5, 2))
        self.status_label.pack(side=tk.LEFT)
        
        self.word_count_label = ttk.Label(status, text="Words: 0", padding=(5, 2))
        self.word_count_label.pack(side=tk.RIGHT)
        
    def setup_bindings(self):
        self.text_area.bind('<KeyRelease>', self.update_word_count)
        self.text_area.bind('<Control-s>', lambda e: self.save_document())
        self.text_area.bind('<Control-o>', lambda e: self.open_document())
        self.text_area.bind('<Control-n>', lambda e: self.new_document())
        self.text_area.bind('<Control-f>', lambda e: self.show_find_dialog())
        
    def new_document(self):
        if self.is_modified:
            if not messagebox.askyesno("Save Changes", 
                "Do you want to save changes to the current document?"):
                return
            self.save_document()
        
        self.text_area.delete('1.0', tk.END)
        self.current_file = None
        self.is_modified = False
        self.status_label.config(text="New Document")
        
    def open_document(self):
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("Word Documents", "*.docx"),
                ("Text Files", "*.txt"),
                ("All Files", "*.*")
            ]
        )
        if file_path:
            self.load_document(file_path)
            
    def load_document(self, file_path):
        try:
            if file_path.endswith('.pdf'):
                # Handle PDF files
                pdf_document = fitz.open(file_path)
                text = ""
                for page in pdf_document:
                    text += page.get_text()
                pdf_document.close()
            elif file_path.endswith('.docx'):
                doc = Document(file_path)
                text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            else:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                    
            self.text_area.delete('1.0', tk.END)
            self.text_area.insert('1.0', text)
            self.current_file = file_path
            self.is_modified = False
            self.status_label.config(text=f"Opened: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open document: {str(e)}")
            
    def save_document(self):
        if not self.current_file:
            return self.save_as_document()
            
        try:
            text = self.text_area.get('1.0', tk.END).strip()
            if not text:
                messagebox.showwarning("Warning", "Document is empty")
                return False
                
            if self.current_file.endswith('.docx'):
                doc = Document()
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():  # Only add non-empty paragraphs
                        doc.add_paragraph(para.strip())
                doc.save(self.current_file)
            else:
                with open(self.current_file, 'w', encoding='utf-8') as file:
                    file.write(text)
                    
            self.is_modified = False
            self.status_label.config(text=f"Saved: {os.path.basename(self.current_file)}")
            return True
            
        except Exception as e:
            error_msg = f"Failed to save document: {str(e)}"
            messagebox.showerror("Error", error_msg)
            self.status_label.config(text="Error saving document")
            return False
            
    def save_as_document(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[
                ("Word Documents", "*.docx"),
                ("Text Files", "*.txt"),
                ("All Files", "*.*")
            ]
        )
        if file_path:
            old_file = self.current_file
            self.current_file = file_path
            if not self.save_document():
                self.current_file = old_file  # Restore old file path if save failed
                return False
            return True
        return False
            
    def show_find_dialog(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Find and Replace")
        dialog.geometry("300x150")
        
        ttk.Label(dialog, text="Find:").grid(row=0, column=0, padx=5, pady=5)
        find_entry = ttk.Entry(dialog)
        find_entry.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(dialog, text="Replace with:").grid(row=1, column=0, padx=5, pady=5)
        replace_entry = ttk.Entry(dialog)
        replace_entry.grid(row=1, column=1, padx=5, pady=5)
        
        def find_text():
            text = find_entry.get()
            start = self.text_area.search(text, '1.0', tk.END)
            if start:
                end = f"{start}+{len(text)}c"
                self.text_area.tag_remove('sel', '1.0', tk.END)
                self.text_area.tag_add('sel', start, end)
                self.text_area.see(start)
                self.text_area.focus()
                
        def replace_text():
            if self.text_area.tag_ranges('sel'):
                self.text_area.delete('sel.first', 'sel.last')
                self.text_area.insert('insert', replace_entry.get())
                
        ttk.Button(dialog, text="Find", command=find_text).grid(row=2, column=0, padx=5, pady=5)
        ttk.Button(dialog, text="Replace", command=replace_text).grid(row=2, column=1, padx=5, pady=5)
        
    def toggle_voice_typing(self):
        try:
            if not self.voice_typing_active:
                # Test microphone availability
                with sr.Microphone() as source:
                    self.status_label.config(text="Initializing microphone...")
                    r = sr.Recognizer()
                    r.adjust_for_ambient_noise(source, duration=1)
                    self.voice_typing_active = True
                    self.voice_button.config(text="Stop Voice Typing")
                    self.status_label.config(text="Listening...")
                    threading.Thread(target=self.voice_typing, daemon=True).start()
            else:
                self.voice_typing_active = False
                self.voice_button.config(text="Voice Typing")
                self.status_label.config(text="Voice typing stopped")
        except (AttributeError, OSError) as e:
            self.voice_button.config(state='disabled')
            messagebox.showerror("Error", 
                "Voice typing is not available.\nMake sure you have a microphone and PyAudio installed.")
            
    def voice_typing(self):
        r = sr.Recognizer()
        try:
            with sr.Microphone() as source:
                r.adjust_for_ambient_noise(source)
                while self.voice_typing_active:
                    try:
                        self.status_label.config(text="Listening...")
                        audio = r.listen(source, timeout=2, phrase_time_limit=10)
                        self.status_label.config(text="Processing speech...")
                        text = r.recognize_google(audio)
                        if text:
                            self.text_area.insert(tk.INSERT, text + " ")
                            self.status_label.config(text="Listening...")
                    except sr.WaitTimeoutError:
                        continue
                    except sr.RequestError:
                        self.status_label.config(text="Network error - voice recognition unavailable")
                        self.voice_typing_active = False
                        self.voice_button.config(text="Voice Typing")
                        break
                    except sr.UnknownValueError:
                        # Speech was unintelligible
                        continue
                    except Exception as e:
                        self.status_label.config(text=f"Error: {str(e)}")
                        time.sleep(1)
                        continue
        except Exception as e:
            self.status_label.config(text=f"Microphone error: {str(e)}")
            self.voice_typing_active = False
            self.voice_button.config(text="Voice Typing")
            
    def toggle_theme(self):
        if self.current_theme == "light":
            self.current_theme = "dark"
            self.text_area.configure(bg="black", fg="white", insertbackground="white")
            self.root.configure(bg="gray20")
        else:
            self.current_theme = "light"
            self.text_area.configure(bg="white", fg="black", insertbackground="black")
            self.root.configure(bg="white")
            
    def apply_font(self):
        font_name = self.font_var.get()
        font_size = int(self.size_var.get())
        
        if self.text_area.tag_ranges("sel"):
            self.text_area.tag_add("font_change", "sel.first", "sel.last")
        else:
            current_pos = self.text_area.index(tk.INSERT)
            self.text_area.tag_add("font_change", current_pos, f"{current_pos}+1c")
            
        self.text_area.tag_configure("font_change", font=(font_name, font_size))
            
    def choose_color(self):
        color = colorchooser.askcolor(title="Choose Text Color")[1]
        if color:
            if self.text_area.tag_ranges("sel"):
                self.text_area.tag_add("color", "sel.first", "sel.last")
                self.text_area.tag_configure("color", foreground=color)
        
    def update_word_count(self, event=None):
        text = self.text_area.get('1.0', tk.END)
        words = len(text.split())
        self.word_count_label.config(text=f"Words: {words}")
        self.is_modified = True
        
    def start_autosave(self):
        def autosave():
            while True:
                time.sleep(self.autosave_interval)
                if self.is_modified and self.current_file:
                    self.save_document()
                    
        threading.Thread(target=autosave, daemon=True).start()
        
    def save_settings(self):
        settings = {
            "theme": self.current_theme,
            "font_family": self.font_var.get(),
            "font_size": self.size_var.get(),
            "autosave_interval": self.autosave_interval
        }
        with open("editor_settings.json", "w") as f:
            json.dump(settings, f)
            
    def load_settings(self):
        try:
            with open("editor_settings.json", "r") as f:
                settings = json.load(f)
                self.current_theme = settings.get("theme", "light")
                self.font_var.set(settings.get("font_family", "Calibri"))
                self.size_var.set(settings.get("font_size", "11"))
                self.autosave_interval = settings.get("autosave_interval", 300)
                self.apply_settings()
        except FileNotFoundError:
            pass
            
    def apply_settings(self):
        if self.current_theme == "dark":
            self.toggle_theme()
        self.apply_font()

    def read_text(self):
        """Read selected text or entire document aloud with improved audio handling"""
        if self.is_reading:
            self.stop_reading()
            return
        
        # Get selected text or all text
        if self.text_area.tag_ranges("sel"):
            text = self.text_area.get("sel.first", "sel.last")
        else:
            text = self.text_area.get("1.0", tk.END)
        
        if not text.strip():
            self.status_label.config(text="No text to read")
            return

        try:
            # Detect language of the text
            detected_lang = detect(text)
            # Map detected language to gTTS language code
            lang_map = {
                'ta': 'ta',  # Tamil
                'en': 'en',  # English
                'hi': 'hi',  # Hindi
                'te': 'te',  # Telugu
                'ml': 'ml',  # Malayalam
                'kn': 'kn'   # Kannada
            }
            tts_lang = lang_map.get(detected_lang, 'en')
            
            # Split text into manageable chunks (3000 chars for non-English, 5000 for English)
            chunk_size = 5000 if tts_lang == 'en' else 3000
            text_chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]
            
            self.status_label.config(text=f"Preparing audio (Language: {tts_lang})...")
            self.is_reading = True
            self.read_button.config(text="Stop Reading", state='disabled')
            
            # Create temp directory for audio chunks
            self.temp_dir = tempfile.mkdtemp()
            self.temp_audio_files = []
            
            # Generate audio in a separate thread
            def generate_audio():
                try:
                    for i, chunk in enumerate(text_chunks):
                        if not self.is_reading:
                            return
                            
                        # Update status for long texts
                        if len(text_chunks) > 1:
                            self.status_label.config(
                                text=f"Preparing audio... ({i+1}/{len(text_chunks)})"
                            )
                            
                        chunk_file = os.path.join(self.temp_dir, f'chunk_{i}.mp3')
                        
                        # Configure TTS based on language
                        tts = gTTS(
                            text=chunk.strip(),
                            lang=tts_lang,
                            slow=tts_lang != 'en'  # Slower speed for non-English text
                        )
                        tts.save(chunk_file)
                        
                        # Verify the audio file was created successfully
                        if not os.path.exists(chunk_file) or os.path.getsize(chunk_file) == 0:
                            raise Exception(f"Failed to generate audio for chunk {i+1}")
                            
                        self.temp_audio_files.append(chunk_file)
                    
                    # Enable button and start playback
                    self.read_button.config(state='normal')
                    self.root.after(0, self.play_audio_chunks)
                    
                except Exception as e:
                    self.root.after(0, lambda: self.handle_error(str(e)))
                    
            threading.Thread(target=generate_audio, daemon=True).start()
                
        except Exception as e:
            self.handle_error(str(e))

    def play_audio_chunks(self):
        """Play audio chunks sequentially with improved handling"""
        try:
            if not self.is_reading or not self.temp_audio_files:
                return
                
            # Load and play first chunk
            self.current_chunk = 0
            pygame.mixer.music.load(self.temp_audio_files[0])
            pygame.mixer.music.set_volume(0.8)  # Slightly reduce volume for clarity
            pygame.mixer.music.play()
            
            # Update UI
            self.read_button.config(text="Stop Reading")
            self.status_label.config(text=f"Reading text (Chunk 1/{len(self.temp_audio_files)})...")
            
            def play_next_chunk():
                if not self.is_reading:
                    return
                    
                if not pygame.mixer.music.get_busy():
                    self.current_chunk += 1
                    if self.current_chunk < len(self.temp_audio_files):
                        # Play next chunk
                        try:
                            pygame.mixer.music.load(self.temp_audio_files[self.current_chunk])
                            pygame.mixer.music.play()
                            self.status_label.config(
                                text=f"Reading text (Chunk {self.current_chunk + 1}/{len(self.temp_audio_files)})..."
                            )
                            self.root.after(100, play_next_chunk)
                        except Exception as e:
                            self.handle_error(f"Error playing chunk {self.current_chunk + 1}: {str(e)}")
                    else:
                        # Finished all chunks
                        self.stop_reading()
                else:
                    self.root.after(100, play_next_chunk)
                    
            self.root.after(100, play_next_chunk)
            
        except Exception as e:
            self.handle_error(str(e))

    def stop_reading(self):
        """Stop reading text aloud and clean up resources"""
        if self.is_reading:
            pygame.mixer.music.stop()
            self.is_reading = False
            self.read_button.config(text="Read Text", state='normal')
            self.status_label.config(text="Reading stopped")
            
            # Clean up temp files
            try:
                if hasattr(self, 'temp_audio_files'):
                    for file in self.temp_audio_files:
                        try:
                            os.unlink(file)
                        except:
                            pass
                    self.temp_audio_files = []
                    
                if hasattr(self, 'temp_dir'):
                    try:
                        os.rmdir(self.temp_dir)
                    except:
                        pass
                    del self.temp_dir
            except Exception:
                pass

    def handle_error(self, error_msg):
        """Handle errors in text-to-speech"""
        self.status_label.config(text=f"Error reading text: {error_msg}")
        self.read_button.config(state='normal')
        self.stop_reading()

    def ocr_tamil_cleaned(self, image_path, page_number):
        """Extract text while removing headers and footers"""
        image = Image.open(image_path)
        raw_text = pytesseract.image_to_string(image, lang='tam+eng')
        
        # Detect language
        detected_lang = detect(raw_text)
        
        # Remove headers and footers
        lines = raw_text.splitlines()
        cleaned_lines = []
        for line in lines:
            if line.strip() and not self.is_header_or_footer(line, page_number):
                cleaned_lines.append(line)
        
        return " ".join(cleaned_lines), detected_lang

    def is_header_or_footer(self, line, page_number):
        """Detect if a line is a header or footer"""
        if re.match(rf"^\s*{page_number}\s*$", line):
            return True
        if len(line.strip()) < 5:
            return True
        return False

    def translate_and_rewrite_text(self, text, target_lang='en', max_retries=3):
        """Translate text to target language"""
        if not text.strip():
            return text

        chunks = [text[i:i+2000] for i in range(0, len(text), 2000)]
        translated_chunks = []
        
        for chunk in chunks:
            if not chunk.strip():
                continue
            
            time.sleep(1)
            
            try:
                translator = GoogleTranslator(source='auto', target=target_lang)
                translated_chunk = translator.translate(chunk)
                translated_chunks.append(translated_chunk if translated_chunk else chunk)
            except Exception as chunk_error:
                print(f"Chunk translation error: {chunk_error}")
                translated_chunks.append(chunk)
            
        if translated_chunks:
            return ' '.join(translated_chunks)
        
        return text

    def show_ocr_dialog(self):
        """Show dialog for OCR processing"""
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("All Supported Files", "*.pdf *.jpg *.jpeg *.png *.tiff *.bmp"),
                ("PDF Files", "*.pdf"),
                ("Image Files", "*.jpg *.jpeg *.png *.tiff *.bmp"),
                ("All Files", "*.*")
            ]
        )
        if file_path:
            # Create settings dialog with compact size
            settings_dialog = tk.Toplevel(self.root)
            settings_dialog.title("OCR Settings")
            settings_dialog.geometry("300x250")  # Reduced size
            
            # Configure grid weights
            settings_dialog.grid_columnconfigure(0, weight=1)
            
            # Language selection with combobox
            lang_frame = ttk.LabelFrame(settings_dialog, text="Target Language")
            lang_frame.grid(row=0, column=0, sticky='ew', padx=5, pady=(5,0))
            
            # Language options
            languages = [
                ("English", "en"),
                ("Tamil", "ta"),
                ("Hindi", "hi"),
                ("Telugu", "te"),
                ("Malayalam", "ml"),
                ("Kannada", "kn")
            ]
            
            lang_var = tk.StringVar(value="en")
            lang_combo = ttk.Combobox(lang_frame, textvariable=lang_var, state='readonly', width=15)
            lang_combo['values'] = [lang[0] for lang in languages]  # Display names
            lang_combo.set("English")  # Set default
            lang_combo.pack(padx=5, pady=2)
            
            # Function to map display name to language code
            def get_lang_code(display_name):
                return next(lang[1] for lang in languages if lang[0] == display_name)
            
            # Page range frame - only show for PDF files
            if file_path.lower().endswith('.pdf'):
                page_frame = ttk.LabelFrame(settings_dialog, text="Page Range")
                page_frame.grid(row=1, column=0, sticky='ew', padx=5, pady=(5,0))
                
                # Get total pages
                pdf = fitz.open(file_path)
                total_pages = len(pdf)
                pdf.close()
                
                ttk.Label(page_frame, 
                         text=f"Total pages: {total_pages}").pack(pady=(2,0))
                
                page_var = tk.StringVar()
                page_entry = ttk.Entry(page_frame, textvariable=page_var)
                page_entry.pack(fill=tk.X, padx=5, pady=2)
                ttk.Label(page_frame, 
                         text="Format: 1-3, 5, 7-9").pack(pady=(0,2))
            
            # Audio options - compact layout
            voice_frame = ttk.LabelFrame(settings_dialog, text="Audio Options")
            voice_frame.grid(row=2, column=0, sticky='ew', padx=5, pady=(5,0))
            
            voice_var = tk.BooleanVar(value=False)
            ttk.Checkbutton(voice_frame, text="Generate audio after OCR", 
                           variable=voice_var).pack(padx=5, pady=2)
            
            # Button frame - aligned at bottom
            button_frame = ttk.Frame(settings_dialog)
            button_frame.grid(row=3, column=0, pady=(5,5))
            
            def start_ocr():
                pages_to_process = None
                if file_path.lower().endswith('.pdf'):
                    # Parse page range
                    page_range = page_var.get().strip()
                    if page_range:
                        try:
                            pages_to_process = set()
                            for part in page_range.split(','):
                                part = part.strip()
                                if '-' in part:
                                    start, end = map(int, part.split('-'))
                                    if start < 1 or end > total_pages or start > end:
                                        raise ValueError(f"Invalid page range: {part}")
                                    pages_to_process.update(range(start, end + 1))
                                else:
                                    page_num = int(part)
                                    if page_num < 1 or page_num > total_pages:
                                        raise ValueError(f"Invalid page number: {page_num}")
                                    pages_to_process.add(page_num)
                        except ValueError as e:
                            messagebox.showerror("Error", str(e))
                            return
                
                # Get language code before destroying dialog
                selected_lang = get_lang_code(lang_combo.get())
                generate_audio = voice_var.get()
                
                settings_dialog.destroy()
                self.ocr_cancel_flag = False  # Reset cancel flag
                
                # Run OCR in a separate thread
                threading.Thread(target=lambda: self.process_file_ocr(
                    file_path, selected_lang, pages_to_process, generate_audio), 
                    daemon=True).start()
            
            ttk.Button(button_frame, text="Start OCR", 
                      command=start_ocr).pack(side=tk.LEFT, padx=2)
            ttk.Button(button_frame, text="Cancel", 
                      command=settings_dialog.destroy).pack(side=tk.LEFT, padx=2)
            
            # Make dialog modal
            settings_dialog.transient(self.root)
            settings_dialog.grab_set()
            settings_dialog.focus_set()

    def process_file_ocr(self, file_path, target_lang='en', pages_to_process=None, read_aloud=False):
        """Process file with OCR and translation"""
        try:
            # Create temporary directory for images
            output_folder = os.path.join(os.path.expanduser("~"), "output_images")
            os.makedirs(output_folder, exist_ok=True)
            
            extracted_text = ""
            self.status_label.config(text="Starting OCR processing...")
            
            # Create cancel button in status bar
            cancel_button = ttk.Button(
                self.status_label.master,
                text="Cancel OCR",
                command=lambda: setattr(self, 'ocr_cancel_flag', True)
            )
            cancel_button.pack(side=tk.RIGHT, padx=5)
            
            try:
                if file_path.lower().endswith('.pdf'):
                    # Process PDF
                    pdf_document = fitz.open(file_path)
                    total_pages = len(pdf_document)
                    
                    # If no pages specified, process all pages
                    if not pages_to_process:
                        pages_to_process = range(1, total_pages + 1)
                    
                    for page_number in sorted(pages_to_process):
                        if self.ocr_cancel_flag:
                            self.status_label.config(text="OCR processing cancelled")
                            break
                            
                        # Update status
                        self.status_label.config(text=f"Processing page {page_number} of {total_pages}...")
                        self.root.update_idletasks()
                        
                        page = pdf_document[page_number - 1]
                        pix = page.get_pixmap()
                        image_path = f"{output_folder}/page_{page_number}.jpg"
                        pix.save(image_path)
                        
                        # Process the image
                        page_text, detected_lang = self.ocr_tamil_cleaned(image_path, page_number)
                        
                        # Translate if needed
                        if detected_lang != target_lang:
                            page_text = self.translate_and_rewrite_text(page_text, target_lang)
                            
                        extracted_text += page_text + "\n"
                        
                        # Clean up temporary image
                        os.remove(image_path)
                        
                    pdf_document.close()
                    
                else:
                    # Process single image file
                    if not self.ocr_cancel_flag:
                        self.status_label.config(text="Processing image...")
                        self.root.update_idletasks()
                        
                        # Process the image
                        page_text, detected_lang = self.ocr_tamil_cleaned(file_path, 1)
                        
                        # Translate if needed
                        if detected_lang != target_lang:
                            page_text = self.translate_and_rewrite_text(page_text, target_lang)
                            
                        extracted_text = page_text
                
                # Insert extracted text into text area if not cancelled
                if not self.ocr_cancel_flag:
                    self.text_area.delete('1.0', tk.END)
                    self.text_area.insert('1.0', extracted_text)
                    self.status_label.config(text="OCR processing complete")
                    
                    # Read the text if requested
                    if read_aloud and extracted_text.strip():
                        self.status_label.config(text="Converting text to speech...")
                        self.root.update_idletasks()
                        
                        # Ask user where to save the audio file
                        audio_file = filedialog.asksaveasfilename(
                            defaultextension=".mp3",
                            filetypes=[("MP3 files", "*.mp3")],
                            title="Save Audio As"
                        )
                        
                        if audio_file:
                            try:
                                # Convert text to speech and save
                                tts = gTTS(text=extracted_text, lang=target_lang)
                                tts.save(audio_file)
                                
                                # Ask if user wants to play the audio now
                                if messagebox.askyesno("Play Audio", 
                                                     "Audio file saved. Would you like to play it now?"):
                                    pygame.mixer.music.load(audio_file)
                                    pygame.mixer.music.play()
                                    
                                    # Update status
                                    self.status_label.config(text="Playing audio...")
                                    
                                    # Wait for audio to finish
                                    while pygame.mixer.music.get_busy():
                                        if self.ocr_cancel_flag:
                                            pygame.mixer.music.stop()
                                            break
                                        time.sleep(0.1)
                                    
                                    pygame.mixer.music.unload()
                                
                                self.status_label.config(text=f"Audio saved to: {os.path.basename(audio_file)}")
                                
                            except Exception as e:
                                messagebox.showerror("Error", f"Failed to save audio: {str(e)}")
                        else:
                            self.status_label.config(text="Audio save cancelled")
            
            finally:
                # Clean up cancel button
                cancel_button.destroy()
                
                # Clean up output folder
                try:
                    os.rmdir(output_folder)
                except:
                    pass
            
        except Exception as e:
            messagebox.showerror("Error", f"OCR processing failed: {str(e)}")
            self.status_label.config(text="OCR processing failed")

def main():
    root = ThemedTk(theme="arc")  # Use a modern theme
    app = DocumentEditor(root)
    root.mainloop()

if __name__ == "__main__":
    main() 