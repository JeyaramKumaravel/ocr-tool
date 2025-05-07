import tkinter as tk
from tkinter import ttk, font, filedialog, messagebox, colorchooser
from tkinter.scrolledtext import ScrolledText
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
import threading
from PIL import Image, ImageTk, ImageOps
from io import BytesIO
import speech_recognition as sr
import time
from ttkthemes import ThemedTk
import fitz  # PyMuPDF for PDF export
import pytesseract
from langdetect import detect
from deep_translator import GoogleTranslator
import re
import cv2
import numpy as np
from gtts import gTTS

# Configure Tesseract OCR path
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

class DocumentEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Editor")
        self.root.geometry("1200x800")
        
        # Document state
        self.current_file = None
        self.is_modified = False
        self.current_theme = "light"
        self.autosave_interval = 300  # 5 minutes
        
        # Then setup UI and other components
        self.setup_ui()
        self.setup_bindings()
        self.start_autosave()
        self.load_settings()
        
    def setup_ui(self):
        # Configure modern style for the editor
        style = ttk.Style()
        style.configure('Editor.TFrame', background='#ffffff')
        style.configure('Ribbon.TFrame', background='#ffffff')
        style.configure('Ribbon.TButton', padding=5)
        style.configure('Tool.TButton', padding=3, width=6)
        style.configure('Group.TLabelframe', background='#ffffff')
        style.configure('Group.TLabelframe.Label', 
                       background='#ffffff', 
                       font=('Segoe UI', 9, 'bold'),
                       foreground='#666666')
        style.configure('Ribbon.TNotebook', background='#ffffff')
        style.configure('Ribbon.TNotebook.Tab', 
                       padding=(12, 6), 
                       font=('Segoe UI', 9))
        
        # Create main container
        self.main_container = ttk.Frame(self.root, style='Editor.TFrame')
        self.main_container.pack(fill=tk.BOTH, expand=True)
        
        # Create ribbon with modern styling
        self.create_ribbon()
        
        # Add separator between ribbon and text area
        ttk.Separator(self.main_container, orient='horizontal').pack(fill=tk.X)
        
        # Main text area
        self.create_text_area()
        
        # Status bar
        self.create_status_bar()
        
    def create_ribbon(self):
        # Create main ribbon frame with Word-like styling
        ribbon_frame = ttk.Frame(self.main_container, style='Ribbon.TFrame')
        ribbon_frame.pack(fill=tk.X, pady=0)
        
        # Create notebook for tabs
        self.ribbon_tabs = ttk.Notebook(ribbon_frame)
        self.ribbon_tabs.pack(fill=tk.X)
        
        # Style for tabs
        style = ttk.Style()
        style.configure('Ribbon.TNotebook', background='#f3f2f1')
        style.configure('Ribbon.TNotebook.Tab', padding=(12, 6), font=('Segoe UI', 9))
        style.map('Ribbon.TNotebook.Tab',
            background=[('selected', '#ffffff'), ('!selected', '#f3f2f1')],
            foreground=[('selected', '#0066cc'), ('!selected', '#444444')])
        
        # Create tabs
        home_tab = ttk.Frame(self.ribbon_tabs, style='Ribbon.TFrame')
        view_tab = ttk.Frame(self.ribbon_tabs, style='Ribbon.TFrame')
        tools_tab = ttk.Frame(self.ribbon_tabs, style='Ribbon.TFrame')
        
        self.ribbon_tabs.add(home_tab, text='Home')
        self.ribbon_tabs.add(view_tab, text='View')
        self.ribbon_tabs.add(tools_tab, text='Tools')
        
        # Add content to tabs
        self.create_home_tab(home_tab)
        self.create_view_tab(view_tab)
        self.create_tools_tab(tools_tab)

    def create_home_tab(self, parent):
        # File group
        file_frame = ttk.LabelFrame(parent, text="File", style='Group.TLabelframe')
        file_frame.pack(side=tk.LEFT, padx=2, pady=1, fill=tk.Y)
        
        # Create a frame for file buttons in a row
        file_buttons_frame = ttk.Frame(file_frame)
        file_buttons_frame.pack(pady=2)
        
        file_buttons = [
            ("New Document", "📄", self.new_document),
            ("Open File", "📂", self.open_document),
            ("Save", "💾", self.save_document),
            ("Save As", "📥", self.save_as_document)
        ]
        
        for tooltip, icon, command in file_buttons:
            btn = ttk.Button(file_buttons_frame, text=icon, command=command, style='Tool.TButton')
            btn.pack(side=tk.LEFT, padx=1)
            self.create_tooltip(btn, tooltip)
        
        # Font group
        font_frame = ttk.LabelFrame(parent, text="Font", style='Group.TLabelframe')
        font_frame.pack(side=tk.LEFT, padx=2, pady=1, fill=tk.Y)
        
        # Font controls in two rows
        font_controls = ttk.Frame(font_frame)
        font_controls.pack(padx=2, pady=2)
        
        # First row - Font family
        ttk.Label(font_controls, text="Font:", font=('Segoe UI', 9)).pack(side=tk.LEFT)
        self.font_var = tk.StringVar(value="Calibri")
        font_combo = ttk.Combobox(font_controls, textvariable=self.font_var, width=15)
        font_combo['values'] = sorted(['Calibri', 'Arial', 'Times New Roman', 'Segoe UI', 'Verdana', 'Consolas'])
        font_combo.pack(side=tk.LEFT, padx=2)
        font_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_font())
        
        # Second row - Font size
        size_frame = ttk.Frame(font_frame)
        size_frame.pack(pady=2)
        
        ttk.Label(size_frame, text="Size:", font=('Segoe UI', 9)).pack(side=tk.LEFT)
        self.size_var = tk.StringVar(value="11")
        size_combo = ttk.Combobox(size_frame, textvariable=self.size_var, width=4)
        size_combo['values'] = [8, 9, 10, 11, 12, 14, 16, 18, 20, 22, 24, 28, 36, 48, 72]
        size_combo.pack(side=tk.LEFT, padx=2)
        size_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_font())

        # Text formatting group
        format_frame = ttk.LabelFrame(parent, text="Formatting", style='Group.TLabelframe')
        format_frame.pack(side=tk.LEFT, padx=2, pady=1, fill=tk.Y)

        format_buttons_frame = ttk.Frame(format_frame)
        format_buttons_frame.pack(pady=2)

        # Bold, Italic, Underline buttons
        self.bold_var = tk.BooleanVar(value=False)
        bold_btn = ttk.Button(format_buttons_frame, text="B", width=3, 
                              command=lambda: self.toggle_format('bold'))
        bold_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(bold_btn, "Bold (Ctrl+B)")

        self.italic_var = tk.BooleanVar(value=False)
        italic_btn = ttk.Button(format_buttons_frame, text="I", width=3,
                               command=lambda: self.toggle_format('italic'))
        italic_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(italic_btn, "Italic (Ctrl+I)")

        self.underline_var = tk.BooleanVar(value=False)
        underline_btn = ttk.Button(format_buttons_frame, text="U", width=3,
                                  command=lambda: self.toggle_format('underline'))
        underline_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(underline_btn, "Underline (Ctrl+U)")

        # Text alignment buttons
        align_frame = ttk.Frame(format_frame)
        align_frame.pack(pady=2)

        align_left_btn = ttk.Button(align_frame, text="⫷", width=3,
                                   command=lambda: self.align_text('left'))
        align_left_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(align_left_btn, "Align Left")

        align_center_btn = ttk.Button(align_frame, text="⫼", width=3,
                                     command=lambda: self.align_text('center'))
        align_center_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(align_center_btn, "Center")

        align_right_btn = ttk.Button(align_frame, text="⫸", width=3,
                                    command=lambda: self.align_text('right'))
        align_right_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(align_right_btn, "Align Right")

        # Text color and highlight buttons
        color_frame = ttk.Frame(format_frame)
        color_frame.pack(pady=2)

        text_color_btn = ttk.Button(color_frame, text="A", width=3,
                                   command=self.choose_text_color)
        text_color_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(text_color_btn, "Text Color")

        highlight_btn = ttk.Button(color_frame, text="H", width=3,
                                  command=self.choose_highlight_color)
        highlight_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(highlight_btn, "Highlight Color")

        # Paragraph group
        paragraph_frame = ttk.LabelFrame(parent, text="Paragraph", style='Group.TLabelframe')
        paragraph_frame.pack(side=tk.LEFT, padx=2, pady=1, fill=tk.Y)

        # Bullet and numbering buttons
        list_frame = ttk.Frame(paragraph_frame)
        list_frame.pack(pady=2)

        bullet_btn = ttk.Button(list_frame, text="•", width=3,
                               command=lambda: self.toggle_list('bullet'))
        bullet_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(bullet_btn, "Bullet List")

        number_btn = ttk.Button(list_frame, text="1.", width=3,
                               command=lambda: self.toggle_list('number'))
        number_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(number_btn, "Numbered List")

        # Indentation buttons
        indent_frame = ttk.Frame(paragraph_frame)
        indent_frame.pack(pady=2)

        decrease_indent_btn = ttk.Button(indent_frame, text="←", width=3,
                                       command=lambda: self.change_indent('decrease'))
        decrease_indent_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(decrease_indent_btn, "Decrease Indent")

        increase_indent_btn = ttk.Button(indent_frame, text="→", width=3,
                                       command=lambda: self.change_indent('increase'))
        increase_indent_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(increase_indent_btn, "Increase Indent")

        # Line spacing dropdown
        spacing_frame = ttk.Frame(paragraph_frame)
        spacing_frame.pack(pady=2)

        ttk.Label(spacing_frame, text="Spacing:").pack(side=tk.LEFT, padx=2)
        self.line_spacing_var = tk.StringVar(value="1.0")
        spacing_combo = ttk.Combobox(spacing_frame, textvariable=self.line_spacing_var, width=4)
        spacing_combo['values'] = ["1.0", "1.15", "1.5", "2.0"]
        spacing_combo.pack(side=tk.LEFT, padx=2)
        spacing_combo.bind('<<ComboboxSelected>>', lambda e: self.apply_line_spacing())

    def create_view_tab(self, parent):
        # View group
        view_frame = ttk.LabelFrame(parent, text="View", style='Group.TLabelframe')
        view_frame.pack(side=tk.LEFT, padx=2, pady=1, fill=tk.Y)
        
        # Create a frame for view buttons in a row
        view_buttons_frame = ttk.Frame(view_frame)
        view_buttons_frame.pack(pady=2)
        
        # Theme toggle
        theme_button = ttk.Button(view_buttons_frame, text="🌓", command=self.toggle_theme, style='Tool.TButton')
        theme_button.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(theme_button, "Toggle Light/Dark Theme")

    def create_tools_tab(self, parent):
        # OCR group
        ocr_frame = ttk.LabelFrame(parent, text="OCR Tools", style='Group.TLabelframe')
        ocr_frame.pack(side=tk.LEFT, padx=2, pady=1, fill=tk.Y)
        
        # OCR buttons frame
        ocr_buttons_frame = ttk.Frame(ocr_frame)
        ocr_buttons_frame.pack(pady=2)
        
        # Language selection
        ttk.Label(ocr_buttons_frame, text="Language:").pack(side=tk.LEFT, padx=2)
        self.ocr_lang_var = tk.StringVar(value="eng")
        lang_combo = ttk.Combobox(ocr_buttons_frame, textvariable=self.ocr_lang_var, width=15)
        lang_combo['values'] = [
            "eng (English)", 
            "tam (Tamil)",
            "hin (Hindi)", 
            "tel (Telugu)",
            "mal (Malayalam)",
            "kan (Kannada)",
            "fra (French)",
            "spa (Spanish)",
            "deu (German)",
            "chi_sim (Chinese Simplified)",
            "jpn (Japanese)"
        ]
        lang_combo.pack(side=tk.LEFT, padx=2)
        
        # OCR from Image button
        ocr_img_btn = ttk.Button(ocr_buttons_frame, text="📷", command=self.ocr_from_image, style='Tool.TButton')
        ocr_img_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(ocr_img_btn, "Extract Text from Image")
        
        # OCR from PDF button
        ocr_pdf_btn = ttk.Button(ocr_buttons_frame, text="📄", command=self.ocr_from_pdf, style='Tool.TButton')
        ocr_pdf_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(ocr_pdf_btn, "Extract Text from PDF")
        
        # Cancel OCR button (initially disabled)
        self.cancel_ocr_btn = ttk.Button(ocr_buttons_frame, text="⨂", command=self.cancel_ocr, 
                                        style='Tool.TButton', state='disabled')
        self.cancel_ocr_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(self.cancel_ocr_btn, "Cancel OCR Process")
        
        # Status label
        self.ocr_status_label = ttk.Label(ocr_frame, text="")
        self.ocr_status_label.pack(pady=2)

        # Translation group
        trans_frame = ttk.LabelFrame(parent, text="Translation", style='Group.TLabelframe')
        trans_frame.pack(side=tk.LEFT, padx=2, pady=1, fill=tk.Y)
        
        # Translation controls frame
        trans_controls = ttk.Frame(trans_frame)
        trans_controls.pack(pady=2)
        
        # Source language selection
        ttk.Label(trans_controls, text="From:").pack(side=tk.LEFT, padx=2)
        self.src_lang_var = tk.StringVar(value="en")
        src_combo = ttk.Combobox(trans_controls, textvariable=self.src_lang_var, width=15)
        src_combo['values'] = [
            "en (English)",
            "ta (Tamil)",
            "hi (Hindi)",
            "te (Telugu)",
            "ml (Malayalam)",
            "kn (Kannada)",
            "fr (French)",
            "es (Spanish)",
            "de (German)",
            "zh (Chinese)",
            "ja (Japanese)"
        ]
        src_combo.pack(side=tk.LEFT, padx=2)
        
        # Target language selection
        ttk.Label(trans_controls, text="To:").pack(side=tk.LEFT, padx=2)
        self.target_lang_var = tk.StringVar(value="en")
        target_combo = ttk.Combobox(trans_controls, textvariable=self.target_lang_var, width=15)
        target_combo['values'] = src_combo['values']  # Use same values as source languages
        target_combo.pack(side=tk.LEFT, padx=2)
        
        # Translation buttons frame
        trans_buttons = ttk.Frame(trans_frame)
        trans_buttons.pack(pady=2)
        
        # Translate button
        translate_btn = ttk.Button(trans_buttons, text="🔄", command=self.translate_text, style='Tool.TButton')
        translate_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(translate_btn, "Translate Selected Text")
        
        # Cancel translation button
        self.cancel_trans_btn = ttk.Button(trans_buttons, text="⨂", command=self.cancel_translation, 
                                          style='Tool.TButton', state='disabled')
        self.cancel_trans_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(self.cancel_trans_btn, "Cancel Translation")
        
        # Translation status label
        self.trans_status_label = ttk.Label(trans_frame, text="")
        self.trans_status_label.pack(pady=2)

        # Text-to-Speech group
        tts_frame = ttk.LabelFrame(parent, text="Text to Speech", style='Group.TLabelframe')
        tts_frame.pack(side=tk.LEFT, padx=2, pady=1, fill=tk.Y)
        
        # TTS controls frame
        tts_controls = ttk.Frame(tts_frame)
        tts_controls.pack(pady=2)
        
        # Language selection
        ttk.Label(tts_controls, text="Language:").pack(side=tk.LEFT, padx=2)
        self.tts_lang_var = tk.StringVar(value="en")
        tts_combo = ttk.Combobox(tts_controls, textvariable=self.tts_lang_var, width=15)
        tts_combo['values'] = [
            "en (English)",
            "ta (Tamil)",
            "hi (Hindi)",
            "te (Telugu)",
            "ml (Malayalam)",
            "kn (Kannada)",
            "fr (French)",
            "es (Spanish)",
            "de (German)",
            "zh (Chinese)",
            "ja (Japanese)"
        ]
        tts_combo.pack(side=tk.LEFT, padx=2)
        
        # TTS buttons frame
        tts_buttons = ttk.Frame(tts_frame)
        tts_buttons.pack(pady=2)
        
        # Convert to Speech button
        tts_btn = ttk.Button(tts_buttons, text="🔊", command=self.text_to_speech, style='Tool.TButton')
        tts_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(tts_btn, "Convert Text to Speech")
        
        # Cancel TTS button
        self.cancel_tts_btn = ttk.Button(tts_buttons, text="⨂", command=self.cancel_tts, 
                                        style='Tool.TButton', state='disabled')
        self.cancel_tts_btn.pack(side=tk.LEFT, padx=1)
        self.create_tooltip(self.cancel_tts_btn, "Cancel Text-to-Speech")
        
        # TTS status label
        self.tts_status_label = ttk.Label(tts_frame, text="")
        self.tts_status_label.pack(pady=2)

    def create_tooltip(self, widget, text):
        """Create a tooltip for a given widget"""
        def show_tooltip(event):
            tooltip = tk.Toplevel()
            tooltip.wm_overrideredirect(True)
            tooltip.wm_geometry(f"+{event.x_root+10}+{event.y_root+10}")
            
            label = ttk.Label(tooltip, text=text, background="#ffffe0", 
                             relief='solid', borderwidth=1)
            label.pack()
            
            def hide_tooltip():
                tooltip.destroy()
            
            widget.tooltip = tooltip
            widget.bind('<Leave>', lambda e: hide_tooltip())
            tooltip.bind('<Leave>', lambda e: hide_tooltip())
        
        widget.bind('<Enter>', show_tooltip)

    def create_text_area(self):
        text_frame = ttk.Frame(self.main_container)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)  # Added margins
        
        self.text_area = ScrolledText(text_frame, wrap=tk.WORD, undo=True)
        self.text_area.pack(fill=tk.BOTH, expand=True)
        
        # Configure Word-like text editor styling
        self.text_area.configure(
            font=("Calibri", 11),  # Word's default font
            bg="#ffffff",
            fg="#000000",
            insertbackground="#000000",
            selectbackground="#b5d7ff",  # Word's selection color
            selectforeground="#000000",
            insertwidth=2,
            spacing1=2,  # Line spacing
            spacing2=2,
            spacing3=2,
            padx=10,
            pady=10
        )
        
    def create_status_bar(self):
        status = ttk.Frame(self.main_container)
        status.pack(fill=tk.X, side=tk.BOTTOM)
        
        # Add top border to status bar
        separator = ttk.Separator(self.main_container, orient='horizontal')
        separator.pack(fill=tk.X, side=tk.BOTTOM)
        
        # Word-like status bar styling
        self.status_label = ttk.Label(status, text="Ready", padding=(10, 2))
        self.status_label.pack(side=tk.LEFT)
        
        self.word_count_label = ttk.Label(status, text="Words: 0", padding=(10, 2))
        self.word_count_label.pack(side=tk.RIGHT, padx=10)
        
    def setup_bindings(self):
        self.text_area.bind('<KeyRelease>', self.update_word_count)
        self.text_area.bind('<Control-s>', lambda e: self.save_document())
        self.text_area.bind('<Control-o>', lambda e: self.open_document())
        self.text_area.bind('<Control-n>', lambda e: self.new_document())
        self.text_area.bind('<Control-f>', lambda e: self.show_find_dialog())
        self.text_area.bind('<Control-b>', lambda e: self.toggle_format('bold'))
        self.text_area.bind('<Control-i>', lambda e: self.toggle_format('italic'))
        self.text_area.bind('<Control-u>', lambda e: self.toggle_format('underline'))
        self.text_area.bind('<Control-l>', lambda e: self.align_text('left'))
        self.text_area.bind('<Control-e>', lambda e: self.align_text('center'))
        self.text_area.bind('<Control-r>', lambda e: self.align_text('right'))
        
    def new_document(self):
        if self.is_modified:
            if not messagebox.askyesno("Save Changes", 
                "Do you want to save changes to the current document?"):
                return
            self.save_document()
        
        self.text_area.delete('1.0', tk.END)
        self.current_file = None
        self.is_modified = False
        self.status_label.config(text="New Document")
        
    def open_document(self):
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("Word Documents", "*.docx"),
                ("Text Files", "*.txt"),
                ("All Files", "*.*")
            ]
        )
        if file_path:
            self.load_document(file_path)
            
    def load_document(self, file_path):
        try:
            if file_path.endswith('.pdf'):
                # Handle PDF files
                pdf_document = fitz.open(file_path)
                text = ""
                for page in pdf_document:
                    text += page.get_text()
                pdf_document.close()
            elif file_path.endswith('.docx'):
                doc = Document(file_path)
                text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            else:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                    
            self.text_area.delete('1.0', tk.END)
            self.text_area.insert('1.0', text)
            self.current_file = file_path
            self.is_modified = False
            self.status_label.config(text=f"Opened: {os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open document: {str(e)}")
            
    def save_document(self):
        if not self.current_file:
            return self.save_as_document()
            
        try:
            text = self.text_area.get('1.0', tk.END).strip()
            if not text:
                messagebox.showwarning("Warning", "Document is empty")
                return False
                
            if self.current_file.endswith('.docx'):
                doc = Document()
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():  # Only add non-empty paragraphs
                        doc.add_paragraph(para.strip())
                doc.save(self.current_file)
            else:
                with open(self.current_file, 'w', encoding='utf-8') as file:
                    file.write(text)
                    
            self.is_modified = False
            self.status_label.config(text=f"Saved: {os.path.basename(self.current_file)}")
            return True
            
        except Exception as e:
            error_msg = f"Failed to save document: {str(e)}"
            messagebox.showerror("Error", error_msg)
            self.status_label.config(text="Error saving document")
            return False
            
    def save_as_document(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[
                ("Word Documents", "*.docx"),
                ("Text Files", "*.txt"),
                ("All Files", "*.*")
            ]
        )
        if file_path:
            old_file = self.current_file
            self.current_file = file_path
            if not self.save_document():
                self.current_file = old_file  # Restore old file path if save failed
                return False
            return True
        return False
            
    def show_find_dialog(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Find and Replace")
        dialog.geometry("300x150")
        
        ttk.Label(dialog, text="Find:").grid(row=0, column=0, padx=5, pady=5)
        find_entry = ttk.Entry(dialog)
        find_entry.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Label(dialog, text="Replace with:").grid(row=1, column=0, padx=5, pady=5)
        replace_entry = ttk.Entry(dialog)
        replace_entry.grid(row=1, column=1, padx=5, pady=5)
        
        def find_text():
            text = find_entry.get()
            start = self.text_area.search(text, '1.0', tk.END)
            if start:
                end = f"{start}+{len(text)}c"
                self.text_area.tag_remove('sel', '1.0', tk.END)
                self.text_area.tag_add('sel', start, end)
                self.text_area.see(start)
                self.text_area.focus()
                
        def replace_text():
            if self.text_area.tag_ranges('sel'):
                self.text_area.delete('sel.first', 'sel.last')
                self.text_area.insert('insert', replace_entry.get())
                
        ttk.Button(dialog, text="Find", command=find_text).grid(row=2, column=0, padx=5, pady=5)
        ttk.Button(dialog, text="Replace", command=replace_text).grid(row=2, column=1, padx=5, pady=5)
        
    def toggle_theme(self):
        if self.current_theme == "light":
            self.current_theme = "dark"
            self.text_area.configure(bg="black", fg="white", insertbackground="white")
            self.root.configure(bg="gray20")
        else:
            self.current_theme = "light"
            self.text_area.configure(bg="white", fg="black", insertbackground="black")
            self.root.configure(bg="white")
            
    def apply_font(self):
        font_name = self.font_var.get()
        font_size = int(self.size_var.get())
        
        self.text_area.configure(font=(font_name, font_size))
            
    def update_word_count(self, event=None):
        text = self.text_area.get('1.0', tk.END)
        words = len(text.split())
        self.word_count_label.config(text=f"Words: {words}")
        self.is_modified = True
        
    def start_autosave(self):
        def autosave():
            while True:
                time.sleep(self.autosave_interval)
                if self.is_modified and self.current_file:
                    self.save_document()
                    
        threading.Thread(target=autosave, daemon=True).start()
        
    def save_settings(self):
        settings = {
            "theme": self.current_theme,
            "font_family": self.font_var.get(),
            "font_size": self.size_var.get(),
            "autosave_interval": self.autosave_interval
        }
        with open("editor_settings.json", "w") as f:
            json.dump(settings, f)
            
    def load_settings(self):
        try:
            with open("editor_settings.json", "r") as f:
                settings = json.load(f)
                self.current_theme = settings.get("theme", "light")
                self.font_var.set(settings.get("font_family", "Calibri"))
                self.size_var.set(settings.get("font_size", "11"))
                self.autosave_interval = settings.get("autosave_interval", 300)
                self.apply_settings()
        except FileNotFoundError:
            pass
            
    def apply_settings(self):
        if self.current_theme == "dark":
            self.toggle_theme()
        self.apply_font()

    def ocr_from_image(self):
        """Handle OCR from image files"""
        file_paths = filedialog.askopenfilenames(
            filetypes=[
                ("Image files", "*.png *.jpg *.jpeg *.bmp *.tiff"),
                ("All files", "*.*")
            ]
        )
        
        if file_paths:
            self.process_ocr_files(file_paths, is_pdf=False)

    def ocr_from_pdf(self):
        """Handle OCR from PDF files with page range selection"""
        file_paths = filedialog.askopenfilenames(
            filetypes=[
                ("PDF files", "*.pdf"),
                ("All files", "*.*")
            ]
        )
        
        if file_paths:
            # Show page range dialog
            dialog = tk.Toplevel(self.root)
            dialog.title("PDF Page Range")
            dialog.geometry("300x150")
            dialog.transient(self.root)
            dialog.grab_set()
            
            # Center the dialog
            dialog.update_idletasks()
            x = self.root.winfo_x() + (self.root.winfo_width() - dialog.winfo_width()) // 2
            y = self.root.winfo_y() + (self.root.winfo_height() - dialog.winfo_height()) // 2
            dialog.geometry(f"+{x}+{y}")
            
            # Create and pack widgets
            ttk.Label(dialog, text="Enter page range (e.g., 1-5 or 1,3,5-7):").pack(pady=10)
            range_entry = ttk.Entry(dialog, width=30)
            range_entry.pack(pady=5)
            range_entry.insert(0, "all")
            
            def validate_and_process():
                range_text = range_entry.get().strip()
                if not range_text or range_text.lower() == 'all':
                    page_ranges = None
                else:
                    try:
                        page_ranges = []
                        parts = range_text.split(',')
                        for part in parts:
                            if '-' in part:
                                start, end = map(int, part.split('-'))
                                if start > end:
                                    raise ValueError("Invalid range")
                                page_ranges.extend(range(start-1, end))
                            else:
                                page_ranges.append(int(part)-1)
                        page_ranges = sorted(set(page_ranges))  # Remove duplicates and sort
                    except ValueError:
                        messagebox.showerror("Error", "Invalid page range format")
                        return
                
                dialog.destroy()
                self.process_ocr_files(file_paths, is_pdf=True, page_ranges=page_ranges)
            
            def show_help():
                help_text = """Page Range Format:
- Use 'all' for all pages
- Single page: 1
- Page range: 1-5
- Multiple ranges: 1-3,5,7-9
- First page is 1"""
                messagebox.showinfo("Page Range Help", help_text)
            
            button_frame = ttk.Frame(dialog)
            button_frame.pack(pady=10)
            
            ttk.Button(button_frame, text="OK", command=validate_and_process).pack(side=tk.LEFT, padx=5)
            ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
            ttk.Button(button_frame, text="Help", command=show_help).pack(side=tk.LEFT, padx=5)
            
            range_entry.focus_set()
            dialog.bind('<Return>', lambda e: validate_and_process())
            dialog.bind('<Escape>', lambda e: dialog.destroy())

    def format_time(self, seconds):
        """Convert seconds to a readable time format (HH:MM:SS)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        seconds = int(seconds % 60)
        
        if hours > 0:
            return f"{hours}h {minutes}m {seconds}s"
        elif minutes > 0:
            return f"{minutes}m {seconds}s"
        else:
            return f"{seconds}s"

    def process_ocr_files(self, file_paths, is_pdf, page_ranges=None):
        """Process OCR for multiple files with progress tracking"""
        self.ocr_cancelled = False
        self.cancel_ocr_btn.config(state='normal')
        total_files = len(file_paths)
        start_time = time.time()
        
        def process_files():
            extracted_text = []
            
            try:
                for i, file_path in enumerate(file_paths, 1):
                    if self.ocr_cancelled:
                        self.update_ocr_status("OCR cancelled")
                        return
                    
                    # Calculate progress and estimated time
                    progress = (i - 1) / total_files
                    elapsed_time = time.time() - start_time
                    if progress > 0:
                        estimated_total = elapsed_time / progress
                        remaining_time = estimated_total - elapsed_time
                        time_str = f" - {self.format_time(remaining_time)} remaining"
                    else:
                        time_str = ""
                    
                    self.update_ocr_status(
                        f"Processing file {i}/{total_files} ({progress*100:.1f}%){time_str}: "
                        f"{os.path.basename(file_path)}"
                    )
                    
                    if is_pdf:
                        # Convert PDF pages to images and process
                        pdf_doc = fitz.open(file_path)
                        
                        # Determine which pages to process
                        if page_ranges is None:
                            pages_to_process = range(pdf_doc.page_count)
                        else:
                            pages_to_process = [p for p in page_ranges if p < pdf_doc.page_count]
                        
                        total_pages = len(pages_to_process)
                        
                        for idx, page_num in enumerate(pages_to_process):
                            if self.ocr_cancelled:
                                break
                            
                            page_progress = (idx / total_pages + (i - 1)) / total_files
                            elapsed = time.time() - start_time
                            if page_progress > 0:
                                est_total = elapsed / page_progress
                                remaining = est_total - elapsed
                                time_str = f" - {self.format_time(remaining)} remaining"
                            else:
                                time_str = ""
                            
                            self.update_ocr_status(
                                f"Processing page {page_num + 1} of selected pages "
                                f"({page_progress*100:.1f}%){time_str}"
                            )
                            
                            page = pdf_doc[page_num]
                            # Increase DPI for better quality
                            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                            
                            # Convert PyMuPDF pixmap to PIL Image
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            
                            # Rotate image if needed
                            width, height = img.size
                            if width > height:
                                img = img.rotate(90, expand=True)
                            
                            text = self.perform_ocr(img)
                            if text.strip():
                                extracted_text.append(text)
                        
                        pdf_doc.close()
                    else:
                        # Process image file directly with enhanced preprocessing
                        img = Image.open(file_path)
                        # Auto-rotate based on EXIF data if available
                        try:
                            img = ImageOps.exif_transpose(img)
                        except:
                            pass
                        text = self.perform_ocr(img)
                        if text.strip():
                            extracted_text.append(text)
                    
                # Insert extracted text into editor
                if extracted_text:
                    final_text = "\n\n".join(extracted_text)
                    self.text_area.insert(tk.END, final_text + "\n")
                    total_time = time.time() - start_time
                    self.update_ocr_status(f"OCR completed in {self.format_time(total_time)}")
                else:
                    self.update_ocr_status("No text was extracted")
                    
            except Exception as e:
                self.update_ocr_status(f"Error during OCR: {str(e)}")
            
            finally:
                self.cancel_ocr_btn.config(state='disabled')
        
        # Start processing in a separate thread
        threading.Thread(target=process_files, daemon=True).start()

    def perform_ocr(self, image):
        """Perform OCR on a single image"""
        try:
            # Get language code from combo box (strip description)
            lang_code = self.ocr_lang_var.get().split()[0]
            
            # Convert to grayscale and preprocess
            img = image.convert('L')
            cv_img = np.array(img)
            cv_img = cv2.normalize(cv_img, None, 0, 255, cv2.NORM_MINMAX)
            
            # Try different preprocessing methods
            results = []
            
            # Method 1: Adaptive Gaussian Thresholding
            try:
                img1 = cv2.adaptiveThreshold(
                    cv_img, 255,
                    cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                    cv2.THRESH_BINARY,
                    11, 2
                )
                text1 = pytesseract.image_to_string(img1, lang=lang_code)
                if text1.strip():
                    results.append(text1)
            except Exception as e:
                self.update_ocr_status(f"Method 1 failed: {str(e)}")
            
            # Method 2: Otsu's thresholding
            try:
                _, img2 = cv2.threshold(cv_img, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                text2 = pytesseract.image_to_string(img2, lang=lang_code)
                if text2.strip():
                    results.append(text2)
            except Exception as e:
                self.update_ocr_status(f"Method 2 failed: {str(e)}")
            
            # Method 3: Original image
            try:
                text3 = pytesseract.image_to_string(cv_img, lang=lang_code)
                if text3.strip():
                    results.append(text3)
            except Exception as e:
                self.update_ocr_status(f"Method 3 failed: {str(e)}")
            
            # Choose the best result (longest text)
            if results:
                return max(results, key=len).strip()
            
            self.update_ocr_status("No text could be extracted from the image")
            return ""
            
        except Exception as e:
            self.update_ocr_status(f"OCR error: {str(e)}")
            return ""

    def post_process_text(self, text):
        """Post-process OCR text to improve accuracy"""
        if not text:
            return text
        
        try:
            # Remove excessive whitespace while preserving paragraphs
            paragraphs = text.split('\n\n')
            cleaned_paragraphs = []
            for para in paragraphs:
                # Join lines that were split in the middle of sentences
                lines = para.split('\n')
                merged_lines = []
                current_line = ""
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                        
                    # If the previous line doesn't end with punctuation, it's probably a continuation
                    if current_line and not current_line[-1] in '.!?":)]}':
                        current_line += ' ' + line
                    else:
                        if current_line:
                            merged_lines.append(current_line)
                        current_line = line
                
                if current_line:
                    merged_lines.append(current_line)
                
                cleaned_para = ' '.join(merged_lines)
                if cleaned_para:
                    cleaned_paragraphs.append(cleaned_para)
            
            text = '\n\n'.join(cleaned_paragraphs)
            
            # Fix common OCR mistakes
            common_fixes = {
                'l': 'I',  # Common confusion between l and I
                '0': 'O',  # Common confusion between 0 and O
                '|': 'I',  # Common confusion with vertical bar
                '1': 'l',  # Common confusion between 1 and l
                'rn': 'm',  # Common confusion between rn and m
                '[': '(',  # Common bracket confusions
                ']': ')',
                '{': '(',
                '}': ')',
            }
            
            for wrong, correct in common_fixes.items():
                text = text.replace(wrong, correct)
            
            # Fix spacing issues
            text = re.sub(r'(?<=\w)\s+(?=[\.,!?])', '', text)  # Remove space before punctuation
            text = re.sub(r'(?<=[\.,!?])\s*(?=\w)', ' ', text)  # Ensure space after punctuation
            
            # Remove likely OCR errors (single characters surrounded by spaces)
            text = re.sub(r'\s+[b-zB-Z]\s+', ' ', text)
            
            # Fix hyphenation
            text = re.sub(r'(?<=\w)[-]\s+', '-', text)
            text = re.sub(r'\s+[-](?=\w)', '-', text)
            
            # Remove non-printable characters while preserving newlines
            text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
            
            return text.strip()
            
        except Exception as e:
            self.update_ocr_status(f"Post-processing error: {str(e)}")
            return text.strip()

    def update_ocr_status(self, message):
        """Update OCR status message"""
        self.ocr_status_label.config(text=message)

    def cancel_ocr(self):
        """Cancel ongoing OCR process"""
        self.ocr_cancelled = True
        self.cancel_ocr_btn.config(state='disabled')
        self.update_ocr_status("Cancelling OCR...")

    def translate_text(self):
        """Translate selected text or entire document"""
        # Get selected text or all text
        if self.text_area.tag_ranges("sel"):
            text = self.text_area.get("sel.first", "sel.last")
            replace_selection = True
        else:
            text = self.text_area.get("1.0", tk.END)
            replace_selection = False
        
        if not text.strip():
            self.update_trans_status("No text to translate")
            return
        
        # Get language codes
        src_lang = self.src_lang_var.get().split()[0]
        target_lang = self.target_lang_var.get().split()[0]
        
        # Enable cancel button
        self.cancel_trans_btn.config(state='normal')
        self.translation_cancelled = False
        
        def translate_chunks():
            try:
                # Split text into chunks
                chunks = self.split_into_chunks(text, 500)
                translated_chunks = []
                start_time = time.time()
                
                # Create translator instance
                translator = GoogleTranslator(source=src_lang, target=target_lang)
                
                for i, chunk in enumerate(chunks, 1):
                    if self.translation_cancelled:
                        self.update_trans_status("Translation cancelled")
                        return
                    
                    # Progress tracking
                    progress = (i - 1) / len(chunks)
                    elapsed_time = time.time() - start_time
                    if progress > 0:
                        estimated_total = elapsed_time / progress
                        remaining_time = estimated_total - elapsed_time
                        time_str = f" - {self.format_time(remaining_time)} remaining"
                    else:
                        time_str = ""
                    
                    self.update_trans_status(
                        f"Translating chunk {i}/{len(chunks)} ({progress*100:.1f}%){time_str}"
                    )
                    
                    try:
                        translated_text = translator.translate(chunk)
                        if translated_text and not translated_text.isspace():
                            translated_chunks.append(translated_text)
                    except Exception as e:
                        self.update_trans_status(f"Error translating chunk {i}: {str(e)}")
                        time.sleep(1)
                
                if translated_chunks:
                    # Post-process the translated text
                    final_text = self.post_process_translation(translated_chunks)
                    
                    # Replace text in editor
                    if replace_selection:
                        self.text_area.delete("sel.first", "sel.last")
                        self.text_area.insert("insert", final_text)
                    else:
                        self.text_area.delete("1.0", tk.END)
                        self.text_area.insert("1.0", final_text)
                    
                    total_time = time.time() - start_time
                    self.update_trans_status(f"Translation completed in {self.format_time(total_time)}")
                else:
                    self.update_trans_status("No text was translated")
                
            except Exception as e:
                self.update_trans_status(f"Translation error: {str(e)}")
            
            finally:
                self.cancel_trans_btn.config(state='disabled')
        
        # Start translation in a separate thread
        threading.Thread(target=translate_chunks, daemon=True).start()

    def post_process_translation(self, translated_chunks):
        """Clean up translated text"""
        text = ' '.join(translated_chunks)
        
        # Basic cleanup
        text = re.sub(r'\s+', ' ', text)  # Remove multiple spaces
        text = re.sub(r'\s*\n\s*', '\n', text)  # Fix newlines
        text = text.strip()
        
        return text

    def split_into_chunks(self, text, chunk_size):
        """Split text into chunks by sentences"""
        chunks = []
        current_chunk = []
        current_size = 0
        
        sentences = text.split('.')
        
        for sentence in sentences:
            sentence = sentence.strip() + '.'
            
            if current_size + len(sentence) > chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_size = len(sentence)
            else:
                current_chunk.append(sentence)
                current_size += len(sentence)
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks

    def update_trans_status(self, message):
        """Update translation status message"""
        self.trans_status_label.config(text=message)

    def cancel_translation(self):
        """Cancel ongoing translation process"""
        self.translation_cancelled = True
        self.cancel_trans_btn.config(state='disabled')
        self.update_trans_status("Cancelling translation...")

    def text_to_speech(self):
        """Convert selected text or entire document to speech"""
        # Get selected text or all text
        if self.text_area.tag_ranges("sel"):
            text = self.text_area.get("sel.first", "sel.last")
        else:
            text = self.text_area.get("1.0", tk.END)
        
        if not text.strip():
            self.tts_status_label.config(text="No text to convert")
            return
            
        # Get language code (strip description)
        lang_code = self.tts_lang_var.get().split()[0]
        
        # Ask user for save location
        if self.current_file:
            base_name = os.path.splitext(os.path.basename(self.current_file))[0]
        else:
            base_name = f"speech_{int(time.time())}"
            
        initial_file = f"{base_name}_{lang_code}.mp3"
        output_file = filedialog.asksaveasfilename(
            defaultextension=".mp3",
            initialfile=initial_file,
            filetypes=[("MP3 files", "*.mp3"), ("All files", "*.*")]
        )
        
        if not output_file:  # User cancelled
            return
            
        # Create output directory if it doesn't exist
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
            except Exception as e:
                messagebox.showerror("Error", f"Could not create directory: {str(e)}")
                return
        
        # Enable cancel button and clear cancel flag
        self.cancel_tts_btn.config(state='normal')
        self.tts_cancelled = False
        
        def convert_to_speech():
            try:
                # Split text into smaller chunks
                chunks = text.split('. ')
                total_chunks = len(chunks)
                
                self.tts_status_label.config(text="Preparing text-to-speech conversion...")
                
                # Initialize timing variables
                start_time = time.time()
                processed_chunks = 0
                
                # Create temp directory in same location as output file
                temp_dir = os.path.join(output_dir, "temp_tts")
                os.makedirs(temp_dir, exist_ok=True)
                
                # Process chunks and show progress
                for i, chunk in enumerate(chunks, 1):
                    if self.tts_cancelled:
                        self.tts_status_label.config(text="Text-to-speech conversion cancelled!")
                        return
                    
                    # Add period back if it was removed during split
                    if not chunk.endswith('.'):
                        chunk += '.'
                    
                    # Create temporary file for this chunk
                    temp_file = os.path.join(temp_dir, f"temp_{i}.mp3")
                    
                    try:
                        tts = gTTS(text=chunk, lang=lang_code)
                        tts.save(temp_file)
                        
                        # Update progress
                        processed_chunks += 1
                        progress = (processed_chunks / total_chunks) * 100
                        
                        # Update status with timing information
                        elapsed_time = int(time.time() - start_time)
                        if processed_chunks > 0:
                            avg_time_per_chunk = elapsed_time / processed_chunks
                            remaining_chunks = total_chunks - processed_chunks
                            estimated_time = int(avg_time_per_chunk * remaining_chunks)
                            time_str = f" - {self.format_time(estimated_time)} remaining"
                        else:
                            time_str = ""
                            
                        self.tts_status_label.config(
                            text=f"Converting to speech: {int(progress)}%{time_str}")
                        
                    except Exception as e:
                        print(f"Error processing chunk {i}: {str(e)}")
                        continue
                
                # Combine all temporary files
                self.tts_status_label.config(text="Finalizing audio file...")
                
                # Combine temporary files into final output
                with open(output_file, 'wb') as outfile:
                    for j in range(1, total_chunks + 1):
                        temp_file = os.path.join(temp_dir, f"temp_{j}.mp3")
                        if os.path.exists(temp_file):
                            with open(temp_file, 'rb') as infile:
                                outfile.write(infile.read())
                            os.remove(temp_file)
                
                # Clean up temp directory
                try:
                    os.rmdir(temp_dir)
                except:
                    pass
                
                if not self.tts_cancelled:
                    os.system(f"start {output_file}")
                    final_time = int(time.time() - start_time)
                    self.tts_status_label.config(
                        text=f"Speech conversion complete! Total time: {self.format_time(final_time)}")
                else:
                    self.tts_status_label.config(text="Speech conversion cancelled!")
                    if os.path.exists(output_file):
                        os.remove(output_file)
                
            except Exception as e:
                self.tts_status_label.config(text=f"Error during conversion: {str(e)}")
                
            finally:
                self.cancel_tts_btn.config(state='disabled')
                # Clean up temp directory if it still exists
                try:
                    if os.path.exists(temp_dir):
                        for f in os.listdir(temp_dir):
                            os.remove(os.path.join(temp_dir, f))
                        os.rmdir(temp_dir)
                except:
                    pass
                
        # Start conversion in a separate thread
        threading.Thread(target=convert_to_speech, daemon=True).start()

    def cancel_tts(self):
        """Cancel ongoing text-to-speech conversion"""
        self.tts_cancelled = True
        self.cancel_tts_btn.config(state='disabled')
        self.tts_status_label.config(text="Cancelling text-to-speech...")

    def toggle_format(self, format_type):
        """Toggle text formatting (bold, italic, underline)"""
        try:
            current_tags = self.text_area.tag_names("sel.first")
            
            if format_type in current_tags:
                self.text_area.tag_remove(format_type, "sel.first", "sel.last")
            else:
                self.text_area.tag_add(format_type, "sel.first", "sel.last")
            
            # Configure tag appearance
            if format_type == 'bold':
                self.text_area.tag_configure('bold', font=('', 0, 'bold'))
                self.bold_var.set(format_type not in current_tags)
            elif format_type == 'italic':
                self.text_area.tag_configure('italic', font=('', 0, 'italic'))
                self.italic_var.set(format_type not in current_tags)
            elif format_type == 'underline':
                self.text_area.tag_configure('underline', underline=True)
                self.underline_var.set(format_type not in current_tags)
            
        except tk.TclError:  # No selection
            messagebox.showinfo("Info", "Please select text to format")

    def align_text(self, alignment):
        """Align text left, center, or right"""
        try:
            # Remove any existing alignment tags
            for tag in ['left', 'center', 'right']:
                self.text_area.tag_remove(tag, "1.0", tk.END)
            
            # Add new alignment
            if self.text_area.tag_ranges("sel"):
                start = self.text_area.index("sel.first linestart")
                end = self.text_area.index("sel.last lineend")
            else:
                start = self.text_area.index("insert linestart")
                end = self.text_area.index("insert lineend")
            
            self.text_area.tag_add(alignment, start, end)
            self.text_area.tag_configure(alignment, justify=alignment)
            
        except Exception as e:
            messagebox.showerror("Error", f"Could not align text: {str(e)}")

    def choose_text_color(self):
        """Open color chooser for text color"""
        try:
            color = colorchooser.askcolor(title="Choose Text Color")[1]
            if color:
                if self.text_area.tag_ranges("sel"):
                    # Create a unique tag for this color
                    tag_name = f"color_{color.replace('#', '')}"
                    self.text_area.tag_add(tag_name, "sel.first", "sel.last")
                    self.text_area.tag_configure(tag_name, foreground=color)
                else:
                    messagebox.showinfo("Info", "Please select text to color")
                
        except Exception as e:
            messagebox.showerror("Error", f"Could not apply color: {str(e)}")

    def choose_highlight_color(self):
        """Open color chooser for text highlighting"""
        try:
            color = colorchooser.askcolor(title="Choose Highlight Color")[1]
            if color:
                if self.text_area.tag_ranges("sel"):
                    # Create a unique tag for this highlight
                    tag_name = f"highlight_{color.replace('#', '')}"
                    self.text_area.tag_add(tag_name, "sel.first", "sel.last")
                    self.text_area.tag_configure(tag_name, background=color)
                else:
                    messagebox.showinfo("Info", "Please select text to highlight")
                
        except Exception as e:
            messagebox.showerror("Error", f"Could not apply highlight: {str(e)}")

    def toggle_list(self, list_type):
        """Toggle bullet or numbered list"""
        try:
            # Get the current line or selected lines
            if self.text_area.tag_ranges("sel"):
                start = self.text_area.index("sel.first linestart")
                end = self.text_area.index("sel.last lineend")
                lines = self.text_area.get(start, end).split('\n')
            else:
                start = self.text_area.index("insert linestart")
                end = self.text_area.index("insert lineend")
                lines = [self.text_area.get(start, end)]

            # Check if lines are already in list format
            first_line = lines[0].lstrip()
            is_list = False
            if list_type == 'bullet' and first_line.startswith('• '):
                is_list = True
            elif list_type == 'number':
                try:
                    num = int(first_line.split('.')[0])
                    if first_line.startswith(f"{num}. "):
                        is_list = True
                except:
                    pass

            # Remove or add list formatting
            new_lines = []
            for i, line in enumerate(lines, 1):
                stripped_line = line.lstrip('• 0123456789. ')
                if is_list:
                    new_lines.append(stripped_line)
                else:
                    prefix = '• ' if list_type == 'bullet' else f"{i}. "
                    new_lines.append(prefix + stripped_line)

            # Replace text
            self.text_area.delete(start, end)
            self.text_area.insert(start, '\n'.join(new_lines))

        except Exception as e:
            messagebox.showerror("Error", f"Could not toggle list: {str(e)}")

    def change_indent(self, direction):
        """Increase or decrease paragraph indentation"""
        try:
            # Get the current line or selected lines
            if self.text_area.tag_ranges("sel"):
                start = self.text_area.index("sel.first linestart")
                end = self.text_area.index("sel.last lineend")
            else:
                start = self.text_area.index("insert linestart")
                end = self.text_area.index("insert lineend")

            # Get current indentation
            current_indent = self.text_area.get(start, end).count('    ')
            
            # Calculate new indentation
            if direction == 'increase':
                new_indent = current_indent + 1
            else:
                new_indent = max(0, current_indent - 1)

            # Apply indentation to each line
            lines = self.text_area.get(start, end).split('\n')
            indented_lines = []
            for line in lines:
                stripped_line = line.lstrip()
                indented_lines.append('    ' * new_indent + stripped_line)

            # Replace text
            self.text_area.delete(start, end)
            self.text_area.insert(start, '\n'.join(indented_lines))

        except Exception as e:
            messagebox.showerror("Error", f"Could not change indentation: {str(e)}")

    def apply_line_spacing(self):
        """Apply line spacing to selected text or entire document"""
        try:
            spacing = float(self.line_spacing_var.get())
            
            # Calculate spacing in pixels (approximate)
            base_spacing = 20  # Base line height in pixels
            pixel_spacing = int(base_spacing * spacing)
            
            # Apply spacing
            self.text_area.configure(spacing1=pixel_spacing//2, 
                                   spacing2=0,
                                   spacing3=pixel_spacing//2)
            
        except Exception as e:
            messagebox.showerror("Error", f"Could not apply line spacing: {str(e)}")

def main():
    root = ThemedTk(theme="arc")  # Use a modern theme
    app = DocumentEditor(root)
    root.mainloop()

if __name__ == "__main__":
    main() 