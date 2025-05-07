import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document
from gtts import gTTS
import os
import re
import time
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from langdetect import detect
from deep_translator import GoogleTranslator
from tkinter import TclError
import tkinterdnd2 as tkdnd

# Configure Tesseract OCR
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
tamil_lang = "tam+eng"  # Tamil language code

# Language Detection and Translation
def translate_and_rewrite_text(text, target_lang='en', max_retries=3):
    if not text.strip():
        return text

    chunks = [text[i:i+2000] for i in range(0, len(text), 2000)]
    translated_chunks = []
    
    for chunk in chunks:
        if not chunk.strip():
            continue
            
        time.sleep(1)
        
        try:
            translator = GoogleTranslator(source='auto', target=target_lang)
            translated_chunk = translator.translate(chunk)
            translated_chunks.append(translated_chunk if translated_chunk else chunk)
        except Exception as chunk_error:
            print(f"Chunk translation error: {chunk_error}")
            translated_chunks.append(chunk)
            
    if translated_chunks:
        return ' '.join(translated_chunks)
        
    return text

# Extract text while removing headers and footers
def ocr_tamil_cleaned(image_path, page_number):
    image = Image.open(image_path)
    raw_text = pytesseract.image_to_string(image, lang=tamil_lang)
    
    # Detect language
    detected_lang = detect(raw_text)
    
    # Optionally print detected language for debugging
    print(f"Detected Language: {detected_lang}")

    # Remove headers and footers
    lines = raw_text.splitlines()
    cleaned_lines = []
    for line in lines:
        if line.strip() and not is_header_or_footer(line, page_number):
            cleaned_lines.append(line)
    
    return " ".join(cleaned_lines), detected_lang

# Detect header/footer (example implementation)
def is_header_or_footer(line, page_number):
    if re.match(rf"^\s*{page_number}\s*$", line):
        return True
    if len(line.strip()) < 5:
        return True
    return False

# Create or append to Word document with PDF name
def create_or_append_word_file_from_pdf(pdf_path, content):
    file_name = os.path.splitext(os.path.basename(pdf_path))[0] + ".docx"
    file_path = os.path.join(os.path.expanduser("~/Documents"), file_name)  # Save to Documents folder
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        doc = Document()  # Create a new document if opening fails
    doc.add_paragraph(content)
    doc.save(file_path)

# Convert text to speech and save as MP3 with PDF name
def text_to_speech_google(pdf_path, text, target_lang='en'):
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
    output_file = os.path.join(os.path.expanduser("~/Documents"), f"{pdf_name}_{target_lang}.mp3")
    
    try:
        # Split text into smaller chunks for progress tracking
        chunks = text.split('. ')
        total_chunks = len(chunks)
        progress_var.set(0)
        
        progress_label.config(text="Preparing text-to-speech conversion...")
        root.update_idletasks()
        
        # Initialize timing variables
        start_time = time.time()
        processed_chunks = 0
        
        # Process chunks and show progress
        for i, chunk in enumerate(chunks, 1):
            if cancel_event.is_set():
                progress_label.config(text="Text-to-speech conversion cancelled!")
                return
                
            # Add period back if it was removed during split
            if not chunk.endswith('.'):
                chunk += '.'
                
            # Create temporary file for this chunk
            temp_file = os.path.join(os.path.expanduser("~/Documents"), f"temp_{i}.mp3")
            
            try:
                tts = gTTS(text=chunk, lang=target_lang)
                tts.save(temp_file)
                
                # Update progress and timing information
                processed_chunks += 1
                progress = (processed_chunks / total_chunks) * 100
                progress_var.set(progress)
                
                # Update elapsed time
                elapsed_time = int(time.time() - start_time)
                elapsed_label.config(text=f"Elapsed Time: {format_time(elapsed_time)}")
                
                # Calculate and update estimated time remaining
                if processed_chunks > 0:
                    avg_time_per_chunk = elapsed_time / processed_chunks
                    remaining_chunks = total_chunks - processed_chunks
                    estimated_time = int(avg_time_per_chunk * remaining_chunks)
                    estimated_label.config(text=f"Estimated Time: {format_time(estimated_time)}")
                
                progress_label.config(text=f"Converting to speech: {int(progress)}%")
                root.update()  # Update the entire GUI
                
            except Exception as e:
                print(f"Error processing chunk {i}: {e}")
                continue
        
        # Combine all temporary files
        progress_label.config(text="Finalizing audio file...")
        root.update_idletasks()
        
        # Use system commands to concatenate MP3 files
        temp_files = [os.path.join(os.path.expanduser("~/Documents"), f"temp_{j}.mp3") 
                     for j in range(1, total_chunks + 1)]
        
        with open(output_file, 'wb') as outfile:
            for temp_file in temp_files:
                if os.path.exists(temp_file):
                    with open(temp_file, 'rb') as infile:
                        outfile.write(infile.read())
                    os.remove(temp_file)  # Clean up temporary file
        
        # Play the final audio file and update final status
        if not cancel_event.is_set():
            os.system(f"start {output_file}")
            final_time = int(time.time() - start_time)
            progress_label.config(text="Speech conversion complete!")
            elapsed_label.config(text=f"Total Time: {format_time(final_time)}")
            estimated_label.config(text="Estimated Time: 0s")
        else:
            progress_label.config(text="Speech conversion cancelled!")
            if os.path.exists(output_file):
                os.remove(output_file)
                
        # Update progress and timing information during conversion
        processed_chunks += 1
        progress = (processed_chunks / total_chunks) * 100
        progress_var.set(progress)
        
        # Update elapsed time
        elapsed_time = int(time.time() - start_time)
        elapsed_label.config(text=f"Elapsed Time: {format_time(elapsed_time)}")
        
        # Calculate and update estimated time remaining
        if processed_chunks > 0:
            avg_time_per_chunk = elapsed_time / processed_chunks
            remaining_chunks = total_chunks - processed_chunks
            estimated_time = int(avg_time_per_chunk * remaining_chunks)
            estimated_label.config(text=f"Estimated Time: {format_time(estimated_time)}")
        
        progress_label.config(text=f"Converting to speech: {int(progress)}%")
        root.update()

    except Exception as e:
        print(f"Error during text-to-speech conversion: {e}")
        progress_label.config(text="Error during conversion. Please check the console for details.")
        # Clean up any temporary files
        for i in range(1, total_chunks + 1):
            temp_file = os.path.join(os.path.expanduser("~/Documents"), f"temp_{i}.mp3")
            if os.path.exists(temp_file):
                os.remove(temp_file)

# Cancellation event
cancel_event = threading.Event()

# Add this helper function near the top of the file
def format_time(seconds):
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    seconds = seconds % 60
    
    if hours > 0:
        return f"{hours}h {minutes}m {seconds}s"
    elif minutes > 0:
        return f"{minutes}m {seconds}s"
    else:
        return f"{seconds}s"

# Process PDF
def process_pdf(pdf_path, pages_to_process, conversion_type, target_lang='en'):
    start_time = time.time()
    output_folder = os.path.join(os.path.expanduser("~"), "output_images")
    
    # Attempt to create the output directory
    try:
        os.makedirs(output_folder, exist_ok=True)
    except PermissionError:
        print(f"Permission denied: Unable to create directory '{output_folder}'. Please check your permissions.")
        return

    # Initialize progress
    progress_var.set(0)
    total_pages = len(pages_to_process)  # Total pages to process based on user selection
    progress_label.config(text=f"Starting PDF processing... ({'All pages' if not page_entry.get().strip() or page_entry.get() == 'e.g., 1-3, 5, 7-9' else 'Specific pages'})")
    elapsed_label.config(text="Elapsed Time: 0s")
    estimated_label.config(text="Estimated Time: --")
    root.update_idletasks()

    try:
        pdf_document = fitz.open(pdf_path)
        extracted_text = ""
        processed_pages = 0

        for page_number in sorted(pages_to_process):  # Process pages in the order specified
            if cancel_event.is_set():
                # Update the UI to show that the task has been canceled
                progress_label.config(text="Task Canceled!")
                elapsed_time = int(time.time() - start_time)
                elapsed_label.config(text=f"Elapsed Time: {format_time(elapsed_time)}")
                estimated_label.config(text="Estimated Time: 0s")
                break  # Exit the loop if canceled

            progress_label.config(text=f"Processing page {page_number} of {total_pages}...")
            root.update_idletasks()  # Update the GUI

            page = pdf_document[page_number - 1]  # Adjust for zero-based index
            pix = page.get_pixmap()
            image_path = f"{output_folder}/page_{page_number}.jpg"
            pix.save(image_path)

            page_text, detected_lang = ocr_tamil_cleaned(image_path, page_number)
            
            if detected_lang != target_lang:
                page_text = translate_and_rewrite_text(page_text, target_lang)

            extracted_text += page_text + "\n"  # Accumulate text for DOCX

            # Update progress
            processed_pages += 1
            progress_var.set((processed_pages / total_pages) * 100)
            progress_label.config(text=f"Completed page {page_number}.")
            
            # Update timing information
            elapsed_time = int(time.time() - start_time)
            elapsed_label.config(text=f"Elapsed Time: {format_time(elapsed_time)}")
            if processed_pages > 0:
                avg_time_per_page = elapsed_time / processed_pages
                remaining_pages = total_pages - processed_pages
                estimated_time = int(avg_time_per_page * remaining_pages)
                estimated_label.config(text=f"Estimated Time: {format_time(estimated_time)}")
            root.update()  # Ensure the GUI updates

        if conversion_type == "DOCX" and not cancel_event.is_set():
            # Include target language in the output file name
            file_name = os.path.splitext(os.path.basename(pdf_path))[0] + f"_{target_lang}.docx"
            create_or_append_word_file_from_pdf(file_name, extracted_text)
            messagebox.showinfo("Success", f"Text saved to {file_name}")

        if conversion_type == "Speech" and not cancel_event.is_set():
            # Include target language in the output file name
            pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
            output_file = os.path.join(os.path.expanduser("~/Documents"), f"{pdf_name}_{target_lang}.mp3")
            text_to_speech_google(output_file, extracted_text, target_lang)

        pdf_document.close()
        if not cancel_event.is_set():
            progress_label.config(text="Processing Complete!")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")
    finally:
        start_button.config(state=tk.NORMAL)
        cancel_button.config(state=tk.DISABLED)
        cancel_event.clear()

# Start processing in a new thread
def start_processing():
    pdf_path = file_entry.get()
    target_lang = target_lang_var.get()

    if not os.path.isfile(pdf_path):
        messagebox.showerror("Error", "Please select a valid PDF file.")
        return

    try:
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)
        pdf_document.close()

        pages_to_process = set()
        page_range_text = page_entry.get().strip()
        
        # Process all pages if entry is empty or contains placeholder
        if page_range_text == "" or page_range_text == "e.g., 1-3, 5, 7-9":
            pages_to_process = set(range(1, total_pages + 1))
        else:
            # Process specified pages
            try:
                for part in page_range_text.split(","):
                    part = part.strip()
                    if "-" in part:
                        start, end = map(int, part.split("-"))
                        if start < 1 or end > total_pages or start > end:
                            raise ValueError(f"Invalid page range: {part}")
                        pages_to_process.update(range(start, end + 1))
                    else:
                        page_num = int(part)
                        if page_num < 1 or page_num > total_pages:
                            raise ValueError(f"Invalid page number: {page_num}")
                        pages_to_process.add(page_num)
            except ValueError as e:
                messagebox.showerror("Error", f"Invalid page range: {str(e)}\nPlease use format like '1-3, 5, 7-9'")
                return

        cancel_event.clear()
        start_button.config(state=tk.DISABLED)
        cancel_button.config(state=tk.NORMAL)
        
        # Run process_pdf in a separate thread
        threading.Thread(
            target=process_pdf, 
            args=(pdf_path, pages_to_process, conversion_var.get(), target_lang), 
            daemon=True
        ).start()

    except Exception as e:
        messagebox.showerror("Error", f"Failed to read PDF: {e}")

# Cancel processing
def cancel_processing():
    cancel_event.set()
    cancel_button.config(state=tk.DISABLED)

# Select PDF file
def select_pdf():
    file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)

        try:
            pdf_document = fitz.open(file_path)
            total_pages_label.config(text=f"Total Pages: {len(pdf_document)}")
            pdf_document.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open PDF: {e}")

# GUI
root = tkdnd.TkinterDnD.Tk()
root.title("OCR and Conversion Tool")
root.grid_rowconfigure(0, weight=1)
root.grid_columnconfigure(0, weight=1)
root.configure(bg="#f5f5f5")  # Light gray background

# Custom styles for ttk widgets
style = ttk.Style()
style.theme_use('clam')  # Use clam theme as base
style.configure('TProgressbar', thickness=20, troughcolor='#E0E0E0', background='#4CAF50')
style.configure('TCombobox', padding=5, background='#FFFFFF')

# Replace container creation and packing with grid
container = tk.Frame(root, bg="#f5f5f5", padx=20, pady=10)
container.grid(row=0, column=0, sticky="nsew")

# Configure container grid
container.grid_columnconfigure(0, weight=1)
for i in range(8):  # For title, subtitle, and 6 main sections
    container.grid_rowconfigure(i, weight=1)

# Update title frame
title_frame = tk.Frame(container, bg="#f5f5f5")
title_frame.grid(row=0, column=0, sticky="ew", pady=(0, 15))
title_frame.grid_columnconfigure(0, weight=1)

title_label = tk.Label(
    title_frame,
    text="OCR Conversion Tool",
    font=("Helvetica", 28, "bold"),
    bg="#f5f5f5",
    fg="#2C3E50"
)
title_label.pack()

subtitle_label = tk.Label(
    title_frame,
    text="Convert PDF documents to text and speech",
    font=("Helvetica", 12),
    bg="#f5f5f5",
    fg="#7F8C8D"
)
subtitle_label.pack()

# Update file frame
file_frame = tk.LabelFrame(
    container,
    text="Document Selection (Browse or Drag & Drop)",
    font=("Helvetica", 12, "bold"),
    bg="#f5f5f5",
    fg="#2C3E50",
    padx=10,
    pady=5
)
file_frame.grid(row=1, column=0, sticky="ew", pady=(0, 10))
file_frame.grid_columnconfigure(1, weight=1)

# Add drop functionality
def handle_drop(event):
    file_path = event.data
    # Remove curly braces if present (Windows)
    file_path = file_path.strip('{}')
    if file_path.lower().endswith('.pdf'):
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        # Update total pages
        try:
            pdf_document = fitz.open(file_path)
            total_pages_label.config(text=f"Total Pages: {len(pdf_document)}")
            pdf_document.close()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open PDF: {e}")
    else:
        messagebox.showerror("Error", "Please drop a PDF file")

# Update file entry with drop binding
file_entry = tk.Entry(
    file_frame,
    font=("Helvetica", 10),
    relief="solid",
    bd=1
)
file_entry.grid(row=0, column=1, sticky="ew", padx=(0, 10))
file_entry.drop_target_register(tkdnd.DND_FILES)
file_entry.dnd_bind('<<Drop>>', handle_drop)

# Add browse button back
browse_button = tk.Button(
    file_frame,
    text="Browse",
    command=select_pdf,
    font=("Helvetica", 10),
    bg="#3498DB",
    fg="white",
    relief="flat",
    padx=15,
    pady=5
)
browse_button.grid(row=0, column=2)

# Add hover effects for browse button
def on_browse_enter(e):
    e.widget['background'] = '#2980B9'

def on_browse_leave(e):
    e.widget['background'] = '#3498DB'

browse_button.bind("<Enter>", on_browse_enter)
browse_button.bind("<Leave>", on_browse_leave)

# Update extraction frame
extraction_frame = tk.LabelFrame(
    container,
    text="Page Range (Optional)",
    font=("Helvetica", 12, "bold"),
    bg="#f5f5f5",
    fg="#2C3E50",
    padx=10,
    pady=5
)
extraction_frame.grid(row=2, column=0, sticky="ew", pady=(0, 10))
extraction_frame.grid_columnconfigure(0, weight=1)

# Remove the combobox and just keep the entry with placeholder
page_entry = tk.Entry(
    extraction_frame,
    font=("Helvetica", 10),
    relief="solid",
    bd=1
)
page_entry.grid(row=0, column=0, sticky="ew")
page_entry.insert(0, "e.g., 1-3, 5, 7-9")  # Add placeholder text
page_entry.config(fg='gray')  # Gray out placeholder text

# Add placeholder behavior
def on_entry_click(event):
    if page_entry.get() == "e.g., 1-3, 5, 7-9":
        page_entry.delete(0, tk.END)
        page_entry.config(fg='black')

def on_focus_out(event):
    if page_entry.get() == "":
        page_entry.insert(0, "e.g., 1-3, 5, 7-9")
        page_entry.config(fg='gray')

page_entry.bind('<FocusIn>', on_entry_click)
page_entry.bind('<FocusOut>', on_focus_out)

# Update language frame
language_frame = tk.LabelFrame(
    container,
    text="Target Language",
    font=("Helvetica", 12, "bold"),
    bg="#f5f5f5",
    fg="#2C3E50",
    padx=10,
    pady=5
)
language_frame.grid(row=3, column=0, sticky="ew", pady=(0, 10))
language_frame.grid_columnconfigure(0, weight=1)

# Target language contents remain the same
target_lang_var = tk.StringVar(value='en')
target_lang_combo = ttk.Combobox(
    language_frame,
    textvariable=target_lang_var,
    values=["en", "ta"],
    state="readonly",
    font=("Helvetica", 10)
)
target_lang_combo.grid(row=0, column=0, sticky="ew")

# Update conversion frame
conversion_frame = tk.LabelFrame(
    container,
    text="Output Format",
    font=("Helvetica", 12, "bold"),
    bg="#f5f5f5",
    fg="#2C3E50",
    padx=10,
    pady=5
)
conversion_frame.grid(row=4, column=0, sticky="ew", pady=(0, 10))
conversion_frame.grid_columnconfigure(0, weight=1)

# Conversion type contents remain the same
conversion_var = tk.StringVar(value="DOCX")
conversion_combo = ttk.Combobox(
    conversion_frame,
    textvariable=conversion_var,
    values=["DOCX", "Speech"],
    state="readonly",
    font=("Helvetica", 10)
)
conversion_combo.grid(row=0, column=0, sticky="ew")

# Update progress section
progress_frame = tk.Frame(container, bg="#f5f5f5")
progress_frame.grid(row=5, column=0, sticky="ew", pady=15)
progress_frame.grid_columnconfigure(0, weight=1)

progress_var = tk.DoubleVar()
progress_bar = ttk.Progressbar(
    progress_frame,
    variable=progress_var,
    maximum=100,
    style='TProgressbar'
)
progress_bar.grid(row=0, column=0, sticky="ew", pady=(0, 10))

# Update status frame
status_frame = tk.Frame(progress_frame, bg="#f5f5f5")
status_frame.grid(row=1, column=0, sticky="ew")
status_frame.grid_columnconfigure(1, weight=1)  # Make middle space expand

progress_label = tk.Label(
    status_frame,
    text="Ready to start",
    font=("Helvetica", 10),
    bg="#f5f5f5",
    fg="#2C3E50"
)
progress_label.grid(row=0, column=0, sticky="w")

total_pages_label = tk.Label(
    status_frame,
    text="Total Pages: --",
    font=("Helvetica", 10),
    bg="#f5f5f5",
    fg="#7F8C8D"
)
total_pages_label.grid(row=0, column=1, sticky="e", padx=(0, 20))

elapsed_label = tk.Label(
    status_frame,
    text="Elapsed Time: 0s",
    font=("Helvetica", 10),
    bg="#f5f5f5",
    fg="#7F8C8D"
)
elapsed_label.grid(row=0, column=2, sticky="e", padx=(0, 20))

estimated_label = tk.Label(
    status_frame,
    text="Estimated Time: --",
    font=("Helvetica", 10),
    bg="#f5f5f5",
    fg="#7F8C8D"
)
estimated_label.grid(row=0, column=3, sticky="e", padx=(0, 20))

# Update button frame
button_frame = tk.Frame(container, bg="#f5f5f5")
button_frame.grid(row=6, column=0, pady=10)

start_button = tk.Button(
    button_frame,
    text="Start",
    command=start_processing,
    font=("Helvetica", 10),
    bg="#2ECC71",
    fg="white",
    relief="flat",
    padx=15,
    pady=5
)
start_button.grid(row=0, column=0, padx=5)

cancel_button = tk.Button(
    button_frame,
    text="Cancel",
    command=cancel_processing,
    font=("Helvetica", 10),
    bg="#E74C3C",
    fg="white",
    relief="flat",
    state=tk.DISABLED,
    padx=15,
    pady=5
)
cancel_button.grid(row=0, column=1, padx=5)

# Add hover effects
def on_enter(e):
    e.widget['background'] = '#27AE60' if e.widget == start_button else '#C0392B'

def on_leave(e):
    e.widget['background'] = '#2ECC71' if e.widget == start_button else '#E74C3C'

start_button.bind("<Enter>", on_enter)
start_button.bind("<Leave>", on_leave)
cancel_button.bind("<Enter>", on_enter)
cancel_button.bind("<Leave>", on_leave)

root.mainloop() 